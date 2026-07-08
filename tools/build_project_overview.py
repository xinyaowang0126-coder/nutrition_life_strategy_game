from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import math
import shutil
import textwrap
from typing import Iterable

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ASSET_ROOT = ROOT / "docs" / "project_overview_assets"
IMAGEGEN_DIR = ASSET_ROOT / "imagegen_series"
UI_DIR = ASSET_ROOT / "ui_reference"
DOC_IMG_DIR = ASSET_ROOT / "doc_images_latest"
OUT_DOCX = ROOT / "docs" / "撑过这一周_项目概述与系统预览.docx"

FONT_REG = Path("C:/Windows/Fonts/NotoSansSC-VF.ttf")
FONT_BOLD = Path("C:/Windows/Fonts/msyhbd.ttc")
FONT_BODY = Path("C:/Windows/Fonts/simsun.ttc")

W, H = 1600, 900


class C:
    bg = "#FFF8F0"
    paper = "#FFFFFF"
    ink = "#2D2D2D"
    muted = "#787878"
    line = "#E4D6C6"
    coral = "#FF6B6B"
    green = "#51CF66"
    yellow = "#FFD43B"
    blue = "#748FFC"
    purple = "#8E7CC3"
    navy = "#2F3A56"
    soft_coral = "#FFE3E3"
    soft_green = "#E5F7E9"
    soft_yellow = "#FFF3BF"
    soft_blue = "#E9EDFF"
    soft_purple = "#F0EAFB"
    shadow = "#D5C3B0"


def font(size: int, bold: bool = False):
    path = FONT_BOLD if bold and FONT_BOLD.exists() else FONT_REG
    if not path.exists():
        path = Path("C:/Windows/Fonts/msyh.ttc")
    return ImageFont.truetype(str(path), size)


F = {
    "title": font(48, True),
    "h1": font(34, True),
    "h2": font(26, True),
    "body": font(21),
    "small": font(17),
    "tiny": font(14),
    "card": font(19, True),
}


def new_canvas(title: str, subtitle: str = ""):
    im = Image.new("RGB", (W, H), C.bg)
    d = ImageDraw.Draw(im)
    for x in range(0, W, 40):
        d.line([(x, 0), (x - 220, H)], fill="#F8EDDF", width=1)
    d.rounded_rectangle((36, 34, W - 36, H - 34), radius=30, fill="#FFFDF8", outline=C.line, width=2)
    d.text((70, 58), title, font=F["title"], fill=C.navy)
    if subtitle:
        d.text((74, 116), subtitle, font=F["body"], fill=C.muted)
    d.line((70, 154, W - 70, 154), fill=C.line, width=2)
    return im, d


def text_size(d: ImageDraw.ImageDraw, s: str, fnt) -> tuple[int, int]:
    box = d.textbbox((0, 0), s, font=fnt)
    return box[2] - box[0], box[3] - box[1]


def wrap_text(d: ImageDraw.ImageDraw, text: str, fnt, max_width: int) -> list[str]:
    lines: list[str] = []
    for para in text.split("\n"):
        if not para:
            lines.append("")
            continue
        current = ""
        for ch in para:
            trial = current + ch
            if text_size(d, trial, fnt)[0] <= max_width:
                current = trial
            else:
                if current:
                    lines.append(current)
                current = ch
        if current:
            lines.append(current)
    return lines


def draw_wrapped(d, xy, text, fnt, fill, max_width, line_gap=8):
    x, y = xy
    for line in wrap_text(d, text, fnt, max_width):
        d.text((x, y), line, font=fnt, fill=fill)
        y += text_size(d, line, fnt)[1] + line_gap
    return y


def panel(d, box, title="", fill=C.paper, outline=C.line, radius=24, title_color=C.navy):
    d.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=2)
    if title:
        d.text((box[0] + 24, box[1] + 18), title, font=F["h2"], fill=title_color)


def pill(d, xy, text, color, fill=None, fnt=None):
    if fnt is None:
        fnt = F["small"]
    fill = fill or "#FFFFFF"
    tw, th = text_size(d, text, fnt)
    x, y = xy
    box = (x, y, x + tw + 28, y + th + 15)
    d.rounded_rectangle(box, radius=18, fill=fill, outline=color, width=2)
    d.text((x + 14, y + 7), text, font=fnt, fill=color)
    return box[2] + 8


def arrow(d, start, end, color=C.muted, width=4):
    d.line((start, end), fill=color, width=width)
    ang = math.atan2(end[1] - start[1], end[0] - start[0])
    size = 14
    p1 = (end[0] - size * math.cos(ang - 0.45), end[1] - size * math.sin(ang - 0.45))
    p2 = (end[0] - size * math.cos(ang + 0.45), end[1] - size * math.sin(ang + 0.45))
    d.polygon([end, p1, p2], fill=color)


def card(d, box, title, body="", accent=C.coral, icon="🍚", fill="#FFFFFF"):
    d.rounded_rectangle(box, radius=22, fill=fill, outline=accent, width=3)
    d.ellipse((box[0] + 18, box[1] + 18, box[0] + 68, box[1] + 68), fill=accent)
    d.text((box[0] + 32, box[1] + 25), icon, font=font(22, True), fill="#FFFFFF")
    d.text((box[0] + 82, box[1] + 20), title, font=F["card"], fill=C.ink)
    if body:
        draw_wrapped(d, (box[0] + 24, box[1] + 84), body, F["small"], C.muted, box[2] - box[0] - 48, 6)


def bar(d, xy, label, value, color):
    x, y = xy
    d.text((x, y), label, font=F["small"], fill=C.ink)
    d.rounded_rectangle((x + 96, y + 5, x + 326, y + 25), radius=10, fill="#EDE2D6")
    d.rounded_rectangle((x + 96, y + 5, x + 96 + int(230 * value), y + 25), radius=10, fill=color)


def save(im: Image.Image, name: str):
    UI_DIR.mkdir(parents=True, exist_ok=True)
    path = UI_DIR / name
    im.save(path)
    return path


def board_style_system():
    im, d = new_canvas("视觉系统与 UI 参考基线", "统一色板、字体层级、组件语言：温暖但不做成营养软件")
    colors = [
        ("背景", C.bg), ("纸面", C.paper), ("文字", C.ink), ("珊瑚", C.coral),
        ("柔绿", C.green), ("金黄", C.yellow), ("蓝紫", C.blue), ("紫色", C.purple),
    ]
    x, y = 95, 205
    for name, col in colors:
        d.rounded_rectangle((x, y, x + 150, y + 90), radius=18, fill=col, outline=C.line)
        d.text((x, y + 104), f"{name} {col}", font=F["small"], fill=C.ink)
        x += 180
    panel(d, (90, 390, 650, 800), "核心 UI 组件", fill="#FFFCF7")
    card(d, (125, 465, 360, 630), "食物卡", "花费、饱腹、心情、饮食负担、标签", C.coral, "🍱")
    card(d, (380, 465, 615, 630), "行动卡", "时间、精力、压力、目标进度", C.green, "📚")
    x = 130
    for label, col in [("确认本餐", C.coral), ("查看详情", C.blue), ("结束当天", C.green)]:
        x = pill(d, (x, 680), label, col, "#FFFFFF", F["body"])
    panel(d, (705, 390, 1510, 800), "字体与语气", fill="#FFFCF7")
    d.text((745, 465), "标题：清楚、温和、有生活感", font=F["h1"], fill=C.navy)
    d.text((745, 525), "正文：解释选择后果，不评价玩家人格。", font=F["body"], fill=C.ink)
    lines = [
        "推荐：这顿帮你撑过了晚上，但身体负担也上来了。",
        "推荐：今天没有很完美，不过节奏还在。",
        "避免：你吃错了。这个不健康。你失败了。",
    ]
    yy = 590
    for line in lines:
        yy = draw_wrapped(d, (745, yy), line, F["body"], C.muted, 690, 10)
    return save(im, "00_style_system.png")


def board_core():
    im, d = new_canvas("项目定位总览", "生活策略 + 卡牌构筑 + 饮食配餐 + 轻度资源管理")
    panel(d, (90, 205, 710, 790), "一句话", fill=C.soft_blue, outline=C.blue)
    draw_wrapped(
        d,
        (130, 285),
        "玩家扮演处在现实生活压力中的角色，在有限天数、预算、时间与精力下，通过饮食、行动和自我调节，尽量撑过这一周。",
        F["h2"],
        C.navy,
        540,
        14,
    )
    panel(d, (760, 205, 1510, 790), "设计原则", fill="#FFFCF7")
    bullets = [
        ("不是营养计算器", "主界面显示生活状态，隐藏复杂营养细项。"),
        ("不是越克制越好", "奶茶、外卖、方便面有短期价值，也有后续代价。"),
        ("稳定比完美重要", "目标是找到可持续节奏，而不是一周全对。"),
        ("角色处境驱动策略", "不同角色的目标、偏好、限制和权重不同。"),
    ]
    yy = 280
    for title, body in bullets:
        d.ellipse((805, yy + 5, 825, yy + 25), fill=C.coral)
        d.text((845, yy), title, font=F["h2"], fill=C.ink)
        yy = draw_wrapped(d, (845, yy + 42), body, F["body"], C.muted, 600, 8) + 22
    return save(im, "01_core_positioning.png")


def board_daily_loop():
    im, d = new_canvas("每日流程", "一局 7 天或 14 天，每天由三餐、行动、睡眠与结算构成")
    steps = [
        ("早晨简报", "剧情 / 今日状态 / 事件", "🌅", C.coral),
        ("早餐", "选择 1-2 张食物卡", "🍚", C.yellow),
        ("午餐", "餐盘组合与即时结算", "🍱", C.green),
        ("晚餐", "饱腹、心情、负担变化", "🍲", C.blue),
        ("自由行动", "复习 / 小睡 / 散步 / 喝水", "📚", C.purple),
        ("睡眠选择", "普通睡 / 早睡 / 后续扩展熬夜", "🌙", C.navy),
        ("每日结算", "压力、预算、饮食质量、稳定度", "📋", C.coral),
    ]
    center = (800, 500)
    radius = 285
    prev = None
    positions = []
    for i, step in enumerate(steps):
        ang = -math.pi / 2 + i * (2 * math.pi / len(steps))
        x = center[0] + radius * math.cos(ang)
        y = center[1] + radius * math.sin(ang)
        positions.append((int(x), int(y)))
    for i, pos in enumerate(positions):
        nxt = positions[(i + 1) % len(positions)]
        arrow(d, (pos[0], pos[1]), (nxt[0], nxt[1]), "#B7A492", 4)
    for (title, body, icon, col), (x, y) in zip(steps, positions):
        box = (x - 135, y - 70, x + 135, y + 80)
        card(d, box, title, body, col, icon, "#FFFCF7")
    d.ellipse((660, 360, 940, 640), fill=C.soft_green, outline=C.green, width=4)
    d.text((700, 430), "第 X 天", font=F["h1"], fill=C.navy)
    d.text((710, 500), "生活节奏", font=F["h2"], fill=C.muted)
    return save(im, "02_daily_loop.png")


def board_main_hud():
    im, d = new_canvas("主界面 UI 参考", "移动竖屏优先，PC 横屏重排；所有文字应清晰、温和、可扫描")
    panel(d, (95, 200, 1495, 805), "", fill="#FDF7ED")
    d.rounded_rectangle((130, 230, 1460, 290), radius=22, fill="#FFFFFF", outline=C.line, width=2)
    d.text((160, 248), "第 3 天 / 7 天    晚上", font=F["h2"], fill=C.navy)
    d.text((1220, 248), "设置", font=F["body"], fill=C.muted)
    d.rounded_rectangle((130, 310, 430, 760), radius=24, fill="#FFFFFF", outline=C.line, width=2)
    d.text((165, 335), "角色状态", font=F["h2"], fill=C.navy)
    for i, (label, val, col) in enumerate([
        ("稳定度", .62, C.coral), ("余额", .45, C.yellow), ("精力", .38, C.green),
        ("心情", .42, C.purple), ("饱腹感", .30, C.blue), ("压力", .72, C.coral),
        ("复习进度", .48, C.navy),
    ]):
        bar(d, (165, 395 + i * 46), label, val, col)
    d.rounded_rectangle((455, 310, 1050, 610), radius=24, fill="#FFFFFF", outline=C.line, width=2)
    d.text((490, 335), "当前餐盘 / 行动区", font=F["h2"], fill=C.navy)
    for i, label in enumerate(["早餐", "午餐", "晚餐"]):
        x = 510 + i * 170
        d.rounded_rectangle((x, 405, x + 140, 520), radius=20, fill=C.bg, outline=[C.coral, C.yellow, C.blue][i], width=3)
        d.text((x + 42, 445), label, font=F["body"], fill=C.ink)
    d.rounded_rectangle((855, 405, 1005, 520), radius=20, fill=C.soft_green, outline=C.green, width=3)
    d.text((882, 445), "组合预览", font=F["body"], fill=C.ink)
    d.rounded_rectangle((1075, 310, 1460, 610), radius=24, fill="#FFFFFF", outline=C.line, width=2)
    d.text((1110, 335), "今日记录", font=F["h2"], fill=C.navy)
    for i, text in enumerate(["早：燕麦 + 牛奶", "午：待选择", "晚：待选择", "事件：考试周", "提示：压力偏高"]):
        d.text((1110, 395 + i * 38), text, font=F["body"], fill=C.ink if i < 3 else C.muted)
    d.rounded_rectangle((455, 635, 1460, 760), radius=24, fill="#FFFFFF", outline=C.line, width=2)
    d.text((490, 655), "手牌区", font=F["h2"], fill=C.navy)
    labels = [("白米饭", "饭", C.coral), ("鸡蛋", "蛋", C.yellow), ("青菜", "菜", C.green), ("奶茶", "茶", C.purple), ("复习", "学", C.blue), ("小睡", "睡", C.navy)]
    x = 590
    for title, icon, col in labels:
        d.rounded_rectangle((x, 655, x + 105, 742), radius=18, fill="#FFFCF7", outline=col, width=2)
        d.ellipse((x + 10, 667, x + 45, 702), fill=col)
        d.text((x + 18, 672), icon, font=F["small"], fill="#FFFFFF")
        d.text((x + 50, 680), title, font=F["small"], fill=C.ink)
        x += 118
    for j, (label, col) in enumerate([("确认本餐", C.coral), ("查看详情", C.blue), ("结束当天", C.green)]):
        d.rounded_rectangle((1275, 648 + j * 34, 1435, 676 + j * 34), radius=14, fill="#FFFFFF", outline=col, width=2)
        d.text((1300, 653 + j * 34), label, font=F["tiny"], fill=col)
    return save(im, "03_main_hud.png")


def board_state_model():
    im, d = new_canvas("稳定度与状态模型", "稳定度不是医学血量，而是“还能不能撑住”的综合生活状态")
    items = [
        ("饮食状态", "餐盘结构 / 蔬果 / 最近三餐", (260, 240), C.green),
        ("心理状态", "心情与压力共同影响", (1030, 240), C.purple),
        ("精力状态", "行动能力与疲劳", (245, 510), C.yellow),
        ("睡眠状态", "次日精力恢复", (1030, 510), C.blue),
        ("饮食负担控制", "油盐糖与暴食压力", (440, 690), C.coral),
        ("目标进度", "复习 / 项目 / 训练", (800, 690), C.navy),
        ("预算安全", "余额能否覆盖剩余天数", (670, 190), C.yellow),
    ]
    for title, body, (x, y), col in items:
        card(d, (x, y, x + 300, y + 130), title, body, col, "●", "#FFFFFF")
        arrow(d, (x + 150, y + 130 if y < 450 else y), (800, 480), col, 3)
    d.ellipse((650, 330, 950, 630), fill=C.soft_coral, outline=C.coral, width=5)
    d.text((715, 420), "稳定度", font=F["h1"], fill=C.navy)
    d.text((735, 485), "0-100", font=F["h2"], fill=C.coral)
    panel(d, (90, 725, 385, 820), "短板惩罚", fill=C.soft_yellow, outline=C.yellow)
    d.text((115, 775), "精力过低 / 压力过高 / 余额危险会额外扣分", font=F["small"], fill=C.ink)
    return save(im, "04_state_model.png")


def board_cards():
    im, d = new_canvas("卡牌系统", "食物、行为、心理卡共同表达选择，而不是单纯健康评分")
    cols = [
        ("食物卡", "花费、饱腹、心情、饮食负担、隐藏营养、偏好加成", C.coral, ["白米饭", "鸡蛋", "青菜", "奶茶", "方便面"]),
        ("行为卡", "时间、精力、压力、目标进度、角色限制", C.green, ["复习", "小睡", "散步", "采购", "运动"]),
        ("心理卡", "自我接纳、压力缓解、习惯锚定、认知重构", C.blue, ["允许不完美", "稳定比极端重要", "今天已经很努力了", "食物不是敌人", "休息也是生产力"]),
    ]
    x = 95
    for title, desc, col, examples in cols:
        panel(d, (x, 220, x + 450, 790), title, fill="#FFFFFF", outline=col)
        draw_wrapped(d, (x + 28, 285), desc, F["body"], C.muted, 390, 8)
        yy = 380
        for ex in examples:
            d.rounded_rectangle((x + 35, yy, x + 415, yy + 62), radius=18, fill=C.bg, outline=col, width=2)
            d.text((x + 58, yy + 17), ex, font=F["body"], fill=C.ink)
            yy += 76
        x += 505
    return save(im, "05_card_system.png")


def board_combos():
    im, d = new_canvas("组合与牌型", "组合爽感来自“把不完美选择放进整体节奏里”")
    combos = [
        ("均衡餐盘", "主食 + 蛋白 + 蔬果\n饱腹、心情、饮食状态提升", C.green),
        ("快速饱腹", "速食救急\n省时但增加饮食负担", C.yellow),
        ("安慰餐", "喜欢的食物 + 稳定配餐\n心情明显恢复", C.purple),
        ("省钱健康餐", "低价食材组成完整一餐\n预算压力下降", C.blue),
        ("外卖补救", "高负担后搭配水果、喝水、散步\n降低后续失控风险", C.coral),
        ("高纤维 / 高蛋白", "服务饱腹、训练恢复与角色策略", C.navy),
    ]
    for i, (title, body, col) in enumerate(combos):
        x = 105 + (i % 3) * 495
        y = 230 + (i // 3) * 250
        card(d, (x, y, x + 430, y + 195), title, body, col, ["🥗", "🍜", "🧋", "🥚", "🍎", "🍗"][i], "#FFFFFF")
    panel(d, (350, 745, 1250, 820), "设计底线：食物不分“好坏阵营”，每张卡都有正面价值和后续代价。", fill=C.soft_green, outline=C.green)
    return save(im, "06_combo_system.png")


def board_economy():
    im, d = new_canvas("经济、采购与库存", "钱不仅买食物，也买便利、时间和心理安慰")
    stages = [
        ("生活费 / 余额", "本周预算、每日最低消费、余额危险", C.yellow),
        ("选择渠道", "食堂 / 便利店 / 外卖 / 买菜 / 轻食", C.coral),
        ("获得卡牌", "食材卡、成品食物卡、救急卡", C.green),
        ("影响未来", "省钱、耗时、耗精力、保质期、库存", C.blue),
    ]
    x = 95
    centers = []
    for title, body, col in stages:
        card(d, (x, 300, x + 310, 500), title, body, col, "￥", "#FFFFFF")
        centers.append((x + 310, 400))
        x += 380
    for i in range(len(centers) - 1):
        arrow(d, centers[i], (centers[i + 1][0] - 310, centers[i + 1][1]), C.muted, 5)
    panel(d, (145, 600, 700, 790), "渠道差异", fill=C.soft_yellow, outline=C.yellow)
    draw_wrapped(d, (180, 665), "自己做饭低成本但耗时间和精力；外卖省时又安慰，但提高饮食负担；轻食状态好但贵，满足感不一定稳定。", F["body"], C.ink, 480, 8)
    panel(d, (815, 600, 1390, 790), "原型简化", fill=C.soft_blue, outline=C.blue)
    draw_wrapped(d, (850, 665), "MVP 阶段先做固定牌池、抽牌重洗、余额不足救急卡；采购和轻库存放到 Phase 2。", F["body"], C.ink, 500, 8)
    return save(im, "07_economy_inventory.png")


def board_characters():
    im, d = new_canvas("角色系统", "角色差异来自处境、目标、限制、偏好和稳定度权重")
    chars = [
        ("熬夜学生", "7 天\n复习进度\n睡眠短板", C.coral),
        ("加班上班族", "14 天\n项目进度\n时间短板", C.navy),
        ("减脂上班族", "14 天\n饮食焦虑\n反弹风险", C.green),
        ("健身大学生", "14 天\n训练恢复\n预算与蛋白", C.blue),
        ("情绪性进食者", "7 天\n情绪波动\n心理卡增强", C.purple),
    ]
    x = 80
    for title, body, col in chars:
        panel(d, (x, 230, x + 285, 775), "", fill="#FFFFFF", outline=col)
        d.ellipse((x + 75, 270, x + 210, 405), fill=col)
        d.text((x + 115, 310), "人", font=F["h1"], fill="#FFFFFF")
        d.text((x + 55, 450), title, font=F["h2"], fill=C.ink)
        draw_wrapped(d, (x + 45, 510), body, F["body"], C.muted, 215, 12)
        x += 305
    return save(im, "08_character_matrix.png")


def board_events():
    im, d = new_canvas("生活事件系统", "隐性 Boss：环境改变选择价值，而不是出现传统敌人")
    events = [
        ("外卖优惠", "外卖降价、出现率上升、自炊吸引力下降", C.coral),
        ("考试周", "压力自然上升，复习要求提高", C.navy),
        ("加班周", "晚间时间减少，外卖和便利店变多", C.purple),
        ("预算紧张", "高价选项更难，廉价食材价值提高", C.yellow),
        ("聚餐邀请", "心情与社交上升，饮食负担可能上升", C.green),
        ("情绪低谷", "甜食出现率上升，心理卡更重要", C.blue),
        ("冰箱清空", "自炊减少，采购价值提高", C.coral),
        ("天气好 / 发薪日", "正面事件也改变策略窗口", C.green),
    ]
    for i, (title, body, col) in enumerate(events):
        x = 90 + (i % 4) * 375
        y = 220 + (i // 4) * 260
        card(d, (x, y, x + 330, y + 205), title, body, col, "!" if i < 6 else "✓", "#FFFFFF")
    return save(im, "09_life_events.png")


def board_endings():
    im, d = new_canvas("结局与长期成长", "结局综合稳定度、目标进度、心理状态、饮食负担与预算")
    endings = [
        ("稳稳撑过", "稳定度高，目标达标，状态较稳", C.green),
        ("勉强撑过", "过程狼狈，但关键目标完成", C.yellow),
        ("吃得健康但心态崩了", "饮食好，但压力和心情失衡", C.purple),
        ("心情不错但身体负担高", "短期舒服，后续仍需调整", C.coral),
        ("预算撑住但状态不好", "省下钱，却消耗精力和生活质量", C.blue),
        ("没撑住", "稳定度过低或关键目标失败，温和重来", C.navy),
    ]
    for i, (title, body, col) in enumerate(endings):
        x = 95 + (i % 3) * 495
        y = 220 + (i // 3) * 215
        card(d, (x, y, x + 430, y + 165), title, body, col, "★", "#FFFFFF")
    panel(d, (260, 690, 1340, 805), "局外成长：解锁角色、新卡牌、新组合、食谱图鉴、14 天模式与挑战模式。", fill=C.soft_green, outline=C.green)
    return save(im, "10_endings_progression.png")


def board_mvp():
    im, d = new_canvas("MVP 范围与阶段路线", "先验证 7 天“熬夜学生”是否好玩，再扩系统")
    phases = [
        ("Phase 0", "项目骨架\nBoot / MainMenu / Autoload", C.navy),
        ("Phase 1", "纯逻辑原型\n状态机 / 公式 / 测试", C.blue),
        ("Phase 2", "可玩 UI 原型\n点卡 / 过天 / 结局", C.green),
        ("Phase 3", "调参与体验\n5 条路线模拟", C.yellow),
        ("Phase 4", "完整系统\n5 角色 / 事件 / 心理卡", C.coral),
        ("Phase 5-6", "表现与发布\n美术音频 / 多平台", C.purple),
    ]
    x = 90
    for i, (title, body, col) in enumerate(phases):
        y = 280 if i % 2 == 0 else 525
        card(d, (x, y, x + 230, y + 165), title, body, col, "●", "#FFFFFF")
        if i < len(phases) - 1:
            arrow(d, (x + 230, y + 82), (x + 300, 402 if i % 2 == 0 else 607), C.muted, 4)
        x += 245
    panel(d, (120, 705, 1480, 810), "MVP 定义：主菜单进入游戏，完成 7 天，每天吃三餐、做行动、睡觉，状态变化可理解，最后出现合理结局。", fill=C.soft_blue, outline=C.blue)
    return save(im, "11_mvp_scope.png")


def board_architecture():
    im, d = new_canvas("Godot 4.7 数据驱动架构", "Resource 配置 + RunState 运行态 + Systems 逻辑 + Control UI")
    rows = [
        [("配置数据", "角色、卡牌、组合、事件、结局都用 Resource 配置", C.blue),
         ("DataRegistry", "加载 .tres，按 ID 查询，启动时校验", C.green)],
        [("RunState", "单局状态：天数、阶段、手牌、指标、历史", C.coral),
         ("系统层", "每日流程、卡牌结算、指标公式、组合检测、结局判定", C.yellow)],
        [("UI 场景", "主界面、卡牌组件、手牌区、餐盘槽、弹窗", C.purple),
         ("验证", "headless 测试、固定路线模拟、日志检查", C.navy)],
    ]
    boxes = []
    for r, row in enumerate(rows):
        for c, (title, body, col) in enumerate(row):
            x = 130 + c * 700
            y = 230 + r * 190
            card(d, (x, y, x + 560, y + 135), title, body, col, "⚙", "#FFFFFF")
            boxes.append((x, y, x + 560, y + 135))
    arrow(d, (690, 298), (830, 298), C.muted, 4)
    arrow(d, (410, 365), (410, 420), C.muted, 4)
    arrow(d, (1110, 365), (1110, 420), C.muted, 4)
    arrow(d, (690, 488), (830, 488), C.muted, 4)
    arrow(d, (410, 555), (410, 610), C.muted, 4)
    arrow(d, (1110, 555), (1110, 610), C.muted, 4)
    return save(im, "12_godot_architecture.png")


def board_platforms():
    im, d = new_canvas("平台、适配与发布", "移动竖屏优先，PC 横屏重排；首发中文，预留多语言")
    panel(d, (90, 230, 530, 760), "屏幕适配", fill=C.soft_blue, outline=C.blue)
    for i, (label, w, h) in enumerate([("手机竖屏", 90, 170), ("平板", 150, 190), ("PC 横屏", 260, 150)]):
        x = 135 + i * 120
        y = 390 if i < 2 else 410
        d.rounded_rectangle((x, y, x + w, y + h), radius=18, fill="#FFFFFF", outline=C.blue, width=3)
        d.text((x, y - 35), label, font=F["small"], fill=C.ink)
    panel(d, (585, 230, 1015, 760), "可访问性", fill=C.soft_green, outline=C.green)
    for i, text in enumerate(["字体大小", "色盲模式", "音频字幕", "柔和提示", "中途保存"]):
        d.text((640, 330 + i * 70), "✓  " + text, font=F["body"], fill=C.ink)
    panel(d, (1070, 230, 1510, 760), "发布准备", fill=C.soft_yellow, outline=C.yellow)
    for i, text in enumerate(["Windows / macOS / Linux", "Android 测试包", "素材授权检查", "商店截图与文案", "健康主题免责声明"]):
        d.text((1125, 330 + i * 70), "□  " + text, font=F["body"], fill=C.ink)
    return save(im, "13_platform_release.png")


def generate_ui_boards() -> list[Path]:
    UI_DIR.mkdir(parents=True, exist_ok=True)
    for p in UI_DIR.glob("*.png"):
        try:
            p.unlink()
        except PermissionError:
            pass
    funcs = [
        board_style_system,
        board_core,
        board_daily_loop,
        board_main_hud,
        board_state_model,
        board_cards,
        board_combos,
        board_economy,
        board_characters,
        board_events,
        board_endings,
        board_mvp,
        board_architecture,
        board_platforms,
    ]
    return [fn() for fn in funcs]


def prepare_doc_images():
    DOC_IMG_DIR.mkdir(parents=True, exist_ok=True)
    for p in DOC_IMG_DIR.glob("*"):
        try:
            p.unlink()
        except PermissionError:
            pass
    for src in sorted(IMAGEGEN_DIR.glob("*.png")):
        img = Image.open(src).convert("RGB")
        img.thumbnail((1500, 844), Image.Resampling.LANCZOS)
        out = DOC_IMG_DIR / (src.stem + ".jpg")
        img.save(out, quality=88, optimize=True)


def set_run_font(run, name="宋体", size=None, color=None, bold=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def set_para_format(p, before=0, after=6, line=1.5, first_indent=False):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if first_indent:
        pf.first_line_indent = Pt(24)


def add_para(doc, text, style="body", bold=False):
    p = doc.add_paragraph()
    if style == "body":
        set_para_format(p, after=7, line=1.35, first_indent=True)
        r = p.add_run(text)
        set_run_font(r, "宋体", 11.5, "2D2D2D", bold)
    elif style == "lead":
        set_para_format(p, before=4, after=10, line=1.25, first_indent=False)
        r = p.add_run(text)
        set_run_font(r, "微软雅黑", 12.5, "2F3A56", bold)
    elif style == "caption":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_format(p, before=2, after=12, line=1.1, first_indent=False)
        r = p.add_run(text)
        set_run_font(r, "宋体", 9.5, "787878", bold)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    set_para_format(p, before=14 if level == 1 else 8, after=6, line=1.15)
    r = p.add_run(text)
    if level == 1:
        set_run_font(r, "微软雅黑", 16, "2E74B5", True)
    else:
        set_run_font(r, "微软雅黑", 13, "2F3A56", True)
    return p


def add_picture(doc, path: Path, caption: str, width=6.35):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_format(p, before=4, after=2, line=1.0)
    p.add_run().add_picture(str(path), width=Inches(width))
    add_para(doc, caption, "caption")


def set_cell_text(cell, text, bold=False, fill=None):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        tc_pr.append(shd)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_format(p, after=0, line=1.15)
    r = p.add_run(text)
    set_run_font(r, "宋体", 10, "2D2D2D", bold)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, True, "F2F4F7")
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    doc.add_paragraph()
    return table


def build_docx(ui_paths: list[Path]):
    prepare_doc_images()
    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_format(p, before=8, after=4, line=1.1)
    r = p.add_run("《撑过这一周》项目概述与系统预览")
    set_run_font(r, "微软雅黑", 20, "2F3A56", True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_format(p, after=12, line=1.15)
    r = p.add_run("生活策略 + 卡牌构筑 + 饮食配餐 + 轻度资源管理")
    set_run_font(r, "微软雅黑", 12.5, "787878", False)
    cover = DOC_IMG_DIR / "01_cover_key_art.jpg"
    if cover.exists():
        add_picture(doc, cover, "图1. 项目主视觉：在不完美现实里寻找可持续节奏", 6.25)
    add_para(doc, "这份文档用于概述《撑过这一周》的核心定位、玩法系统、MVP 范围、视觉基调和 Godot 实现路线。文档中的插画来自 imagegen 统一风格生成；所有 UI 与系统参考图中的中文文字由本地确定性排版生成，可作为后续 UI 设计与实现参考。", "lead", True)

    add_heading(doc, "一、项目定位", 1)
    add_para(doc, "《撑过这一周》不是营养记录工具，也不是“健康食物得高分”的配餐题。玩家扮演一个正在经历现实生活压力的人，在有限天数、预算、时间、精力、情绪和饱腹感之间做取舍，目标是尽量维持稳定度，撑过这一段生活。")
    add_para(doc, "游戏的表达重点是“稳定、可持续、还能继续生活”。奶茶、炸鸡、方便面、外卖并非错误选项，它们可以带来安慰、省时和短期支撑；健康食物也不总是最优，因为它可能更贵、更耗精力，甚至增加压力。")
    add_picture(doc, ui_paths[1], "图2. 项目定位参考板：生活策略而非营养计算器", 6.35)

    add_heading(doc, "二、核心体验与每日循环", 1)
    add_para(doc, "一局游戏以 7 天或 14 天为单位。每天包含早晨简报、三餐选择、自由行动、睡眠选择和每日结算。玩家在每个阶段打出食物卡、行为卡或心理卡，系统即时更新状态，并在晚间汇总饮食质量、压力、预算与目标进度。")
    add_picture(doc, ui_paths[2], "图3. 每日流程参考板：早晨、三餐、行动、睡眠与结算", 6.35)
    add_picture(doc, ui_paths[3], "图4. 主界面 UI 参考板：指标、餐盘、手牌、今日记录与操作区", 6.35)

    add_heading(doc, "三、状态与稳定度系统", 1)
    add_para(doc, "稳定度是项目最重要的表层指标，表示角色“还能不能撑住”。它由饮食状态、心理状态、精力、睡眠、饮食负担控制、目标进度与预算安全共同决定。低精力、低心情、极度饥饿、高压力或经济危机还会触发短板惩罚，避免平均值掩盖关键风险。")
    add_picture(doc, ui_paths[4], "图5. 稳定度模型参考板：多指标共同构成整体生活状态", 6.35)

    add_heading(doc, "四、卡牌与组合系统", 1)
    add_para(doc, "卡牌是玩家行动和饮食选择的主要载体。食物卡负责餐盘与饮食后果，行为卡负责学习、工作、休息、采购和运动，心理卡负责自我接纳、压力缓解和可持续节奏。组合系统提供构筑爽感，让玩家通过搭配获得额外收益。")
    add_picture(doc, ui_paths[5], "图6. 卡牌系统参考板：食物卡、行为卡、心理卡", 6.35)
    add_picture(doc, ui_paths[6], "图7. 组合系统参考板：均衡餐盘、安慰餐、省钱健康餐与外卖补救", 6.35)
    add_table(
        doc,
        ["组合", "触发条件", "设计目的"],
        [
            ["均衡餐盘", "主食 + 蛋白 + 蔬果", "提升饱腹、饮食状态和稳定度"],
            ["快速饱腹", "速食或方便食品", "救急省时，但增加饮食负担"],
            ["安慰餐", "喜欢的食物 + 稳定配餐", "承认安慰价值，降低心理负担"],
            ["外卖补救", "高负担后搭配水果、喝水或散步", "不抹除代价，但降低后续失控风险"],
        ],
        [1.4, 2.2, 2.6],
    )

    add_heading(doc, "五、经济、采购与生活事件", 1)
    add_para(doc, "经济系统用于制造真实限制。钱不仅买食物，也买便利、时间和心理安慰。采购与库存不应过重，MVP 阶段可以先用固定牌池、抽牌重洗和救急卡保证流程不断，Phase 2 再加入食材采购、库存和保质期。")
    add_picture(doc, ui_paths[7], "图8. 经济与库存参考板：生活费、渠道选择、获得卡牌与未来影响", 6.35)
    add_para(doc, "生活事件相当于隐性 Boss。它们不是敌人，而是环境条件变化，例如外卖优惠、考试周、加班周、预算紧张、聚餐邀请、情绪低谷和冰箱清空。这些事件会改变资源压力和选择价值。")
    add_picture(doc, ui_paths[9], "图9. 生活事件参考板：通过环境变化制造挑战", 6.35)

    add_heading(doc, "六、角色、结局与长期成长", 1)
    add_para(doc, "角色差异应体现在能做什么、不能做什么、同一行动的代价、食物偏好、专属目标和稳定度权重上。首个 MVP 聚焦“熬夜学生”，后续扩展加班上班族、减脂上班族、健身大学生和情绪性进食者。")
    add_picture(doc, ui_paths[8], "图10. 角色系统参考板：五个基础生活处境与专属矛盾", 6.35)
    add_para(doc, "结局不只看单一分数，而是综合稳定度、目标进度、心理状态、饮食负担和预算。结局文案必须温和，即便没撑住也应保留重新尝试的空间。局外成长用于解锁角色、卡牌、组合、图鉴、14 天模式和挑战模式。")
    add_picture(doc, ui_paths[10], "图11. 结局与成长参考板：多结局和可重复游玩动力", 6.35)

    add_heading(doc, "七、MVP 与实现路线", 1)
    add_para(doc, "第一优先级是完成 7 天“熬夜学生”可玩原型。MVP 包含 15 张食物卡、4 张行动卡、5 个主状态、2 个半隐藏状态、复习进度、均衡餐盘与快速饱腹两个组合，以及 3 种基础结局。")
    add_picture(doc, ui_paths[11], "图12. MVP 与阶段路线参考板：先验证核心循环，再扩内容", 6.35)
    add_table(
        doc,
        ["阶段", "目标", "验收重点"],
        [
            ["Phase 0", "Godot 项目骨架", "能启动到主菜单，中文字体正常"],
            ["Phase 1", "纯逻辑原型", "固定路线能跑完 7 天，稳定度无异常"],
            ["Phase 2", "可玩 UI 原型", "玩家能点卡、过天、看到结局"],
            ["Phase 3", "调参与体验验证", "放纵、均衡、省钱、硬卷等路线有差异"],
            ["Phase 4+", "完整核心游戏", "多角色、事件、心理卡、图鉴与存档"],
        ],
        [1.2, 2.1, 3.1],
    )

    add_heading(doc, "八、Godot 技术架构", 1)
    add_para(doc, "项目建议采用 Godot 4.7、GDScript、Control/Container UI、Resource 数据配置、JSON 存档、Autoload 少量全局服务，以及系统层集中处理公式和流程。RunState 应作为 Resource 或由 RunController 持有，避免核心运行态使用 RefCounted 引发生命周期问题。")
    add_picture(doc, ui_paths[12], "图13. Godot 数据驱动架构参考板", 6.35)
    add_picture(doc, ui_paths[13], "图14. 平台适配与发布参考板", 6.35)

    add_heading(doc, "九、Imagegen 视觉预览组", 1)
    add_para(doc, "以下为 imagegen 生成的统一风格视觉预览，适合用作情绪板、系统氛围和后续美术方向参考。具体 UI 文字与布局以本文前面的中文可读参考板为准。")
    gallery = [
        ("02_daily_loop_illustration.jpg", "每日循环氛围图"),
        ("04_stability_system_illustration.jpg", "稳定度系统氛围图"),
        ("05_card_taxonomy_illustration.jpg", "卡牌系统氛围图"),
        ("06_combo_system_illustration.jpg", "组合系统氛围图"),
        ("07_economy_inventory_illustration.jpg", "经济库存氛围图"),
        ("08_character_system_illustration.jpg", "角色系统氛围图"),
        ("09_life_events_illustration.jpg", "生活事件氛围图"),
        ("12_godot_architecture_illustration.jpg", "技术架构氛围图"),
        ("13_psychological_tone_illustration.jpg", "心理健康语气氛围图"),
        ("14_platform_release_illustration.jpg", "跨平台发布氛围图"),
    ]
    for idx, (name, cap) in enumerate(gallery, 15):
        p = DOC_IMG_DIR / name
        if p.exists():
            add_picture(doc, p, f"图{idx}. {cap}", 6.1)

    add_heading(doc, "十、后续开发提醒", 1)
    reminders = [
        "先做 headless 数值模拟，再做完整 UI；数值问题越早发现越便宜。",
        "主界面不常驻显示热量、蛋白质、钠、糖等营养细项，避免变成健康记录软件。",
        "每次新增系统都要问：它是否让玩家更能感到“我在过一周生活”？",
        "所有失败、崩盘、失控反馈都要温和，不羞辱玩家。",
        "图片生成可用于情绪和风格探索；UI 文字、按钮、指标名必须确定性排版。",
    ]
    for item in reminders:
        p = doc.add_paragraph(style=None)
        set_para_format(p, after=5, line=1.25, first_indent=False)
        r = p.add_run("• ")
        set_run_font(r, "微软雅黑", 11.5, "FF6B6B", True)
        r = p.add_run(item)
        set_run_font(r, "宋体", 11.5, "2D2D2D", False)

    doc.save(OUT_DOCX)
    return OUT_DOCX


def main():
    ui_paths = generate_ui_boards()
    build_docx(ui_paths)
    print(f"Generated {len(ui_paths)} UI reference boards")
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
